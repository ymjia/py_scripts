# -*- coding: utf-8 -*-
## @file tex_check.py
## @brief get tex information from pdf 
## @author jiayanming

import os.path
import cv2
import numpy as np
import re
import math
import PIL
from PIL import Image
import openpyxl
import docx
from pdf2image import convert_from_path
import tesserocr as to
from tesserocr import PyTessBaseAPI
from operator import itemgetter
from docx.shared import Mm
# docx cell color
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import time

# main: process all xlsx file in current dir
cwd = os.getcwd()
str_input = os.path.join(cwd, "input")
str_output = os.path.join(cwd, "output")

#constant
port_list = {}


class img_name:
    def __init__(self, full_path, pdf, idx):
        self.path = full_path
        self.pdf = pdf
        self.idx = idx

## @brief tex info item from pdf
class tex_item:
    def __init__(self):
        self.date = ""
        self.no = ""
        self.idx = ""
        self.tex_type = ""
        self.port = ""
        self.amount = 0
        self.con_id = ""
        self.img_date = None
        self.img_no = None
        self.img_idx = None
        self.img_tex_type = None
        self.img_port = None
        self.img_amount = None
        self.img_con_id = None
        self.pdf = ""
        self.pdf_idx = ""

## @brief ocr table info (cell pixel positions) from image
class ocr_table:
    def __init__(self, w, h):
        self._h = h # height
        self._w = w # width
        self._origin = (0, 0) # table l-t-coner in image
        self._pos = [] # l-t-coner of each cell
    
    def r(self, rid):
        start = rid * self._w
        return self._pos[start:start + self._h]

    def rect(self, rid, cid):
        rn = rid + 1
        cn = cid + 1
        if rn >= self._h or cn >= self._w:
            return (0, 0, 3, 3)
        top_right = self._pos[rid * self._w + cid]
        low_left = self._pos[rn * self._w + cn]
        if low_left[0] - top_right[0] <= 0 or low_left[1] - top_right[1] <= 0:
            return (0, 0, 3, 3)
        return (top_right[0] + self._origin[0], top_right[1] + self._origin[1],
                low_left[0] + self._origin[0], low_left[1] + self._origin[1])

    def lrect(self, rid, cid):
        rn = rid + 1
        cn = cid + 1
        if rn >= self._h or cn >= self._w:
            return (0, 0, 0, 0)
        top_right = self._pos[rid * self._w + cid]
        low_left = self._pos[rn * self._w + cn]
        return (top_right[0], top_right[1], low_left[0], low_left[1])


## @brief get mean poisition from cv coutour boundingRect
def tuple_mean(tu4):
    return (int(tu4[0] + tu4[2] / 2), int(tu4[1] + tu4[3] / 2))


## @brief generate table information from ocr region
def generate_table(origin, horizontal, vertical, h_pattern, v_pattern):
    dot_pattern = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))

    horizontal = cv2.dilate(horizontal, h_pattern, iterations = 5)
    horizontal = cv2.dilate(horizontal, dot_pattern, iterations = 1)
    horizontal = cv2.erode(horizontal, dot_pattern, iterations = 1)

    vertical = cv2.dilate(vertical, v_pattern, iterations = 5)
    vertical = cv2.dilate(vertical, dot_pattern, iterations = 1)
    vertical = cv2.erode(vertical, dot_pattern, iterations = 1)

    joints = cv2.bitwise_and(horizontal, vertical)
    joints = cv2.dilate(joints, dot_pattern, iterations = 1)
    joints = cv2.erode(joints, dot_pattern, iterations = 1)

    c_dots = cv2.findContours(joints, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)[-2]
    tmp_table = cv2.bitwise_or(horizontal, vertical)

    c_h_lines = cv2.findContours(horizontal, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)[-2]
    c_v_lines = cv2.findContours(vertical, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)[-2]
    d_num = len(c_dots)
    h_num = len(c_h_lines)
    v_num = len(c_v_lines)
    t = ocr_table(v_num, h_num)
    t._origin = origin
    all_dots = []
    for d in c_dots:
        t._pos.append(tuple_mean(cv2.boundingRect(d)))
    # sort by row then col
    # ALTERNATIVE: calulate by v/h line coord 
    t._pos.sort(key = itemgetter(1))
    for li in range(0, h_num):
        start = li * v_num
        end = start + v_num
        sub_line = t._pos[start:end]
        t._pos[start:end] = sorted(sub_line, key = itemgetter(0))
    return t


## @brief detect tables from image file
def ocr_detect_table(iname):
    # load image
    img_path = iname.path
    src = cv2.imread(img_path)
    if src is None:
        print("Error! cannot load image file {}".format(img_path))
    # grey
    grey = cv2.cvtColor(src, cv2.COLOR_BGR2GRAY)
    bw = cv2.adaptiveThreshold(cv2.bitwise_not(grey), 255,
                               cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 51, -5)
    cv2.imwrite(os.path.join(str_output, "bw{}_{}.png".format(iname.pdf, iname.idx)), bw)
    # extract horizontal
    horizontal = bw
    vertical = bw
    scale = 20 # min length to be detected threshold
    h, w = horizontal.shape[:2]
    # Specify size
    horizontalsize = int(w / scale)
    verticalsize = int(h / scale)
    print("{}-{}".format(horizontalsize, verticalsize))
    # Create structure element for extracting lines through morphology operations
    h_pattern = cv2.getStructuringElement(cv2.MORPH_RECT, (horizontalsize, 1))
    v_pattern = cv2.getStructuringElement(cv2.MORPH_RECT, (1, verticalsize))
    # Apply morphology operations
    horizontal = cv2.erode(horizontal, h_pattern, iterations = 1)
    horizontal = cv2.dilate(horizontal, h_pattern, iterations = 1)
    vertical = cv2.erode(vertical, v_pattern)
    vertical = cv2.dilate(vertical, v_pattern)
    joints = cv2.bitwise_and(horizontal, vertical)

    tmp_table = cv2.bitwise_or(horizontal, vertical)

    #Image.fromarray(joints).save(os.path.join(str_output, "{}_{}_joints.png".format(iname.pdf, iname.idx)))
    Image.fromarray(tmp_table).save(os.path.join(str_output, "{}_{}_table.png".format(iname.pdf, iname.idx)))

    # divide and sort tables
    tables = cv2.bitwise_or(horizontal, vertical)
    c_tables = cv2.findContours(tables, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)[-2]
    
    all_tables = []
    for ct in c_tables:
        all_tables.append(cv2.boundingRect(ct))
    all_tables.sort(key=itemgetter(0,1))


    list_table = []
    for tb in all_tables:
        x = tb[0]
        y = tb[1]
        w = tb[2]
        h = tb[3]
        sub_h = horizontal[y:y+h, x:x+w]
        sub_v = vertical[y:y+h, x:x+w]
        list_table.append(generate_table((x, y), sub_h, sub_v, h_pattern, v_pattern))

    return list_table


def my_decode(str_in):
    return str_in.encode('utf-8').decode("gbk", 'ignore')


def regulate_date_str(in_str):
    m = re.search('(20[\d]{2})年([\d]{2})月([\d]{2}).*', in_str)
    if m:
        return "{}/{}/{}".format(m.group(2), m.group(3), m.group(1))
    return ""


def regulate_type_str(in_str):
    if len(in_str) >= 1:
        m = in_str[0]
        if m == "L" or m == "A":
            return m
    return ""


def read_port_list(filename, port_list):
    port_file = filename
    table_port = openpyxl.load_workbook(port_file)
    ws = table_port.active
    for r in ws.iter_rows(min_row=2, max_col=3, values_only=True):
        pid = r[0]
        pname = r[1]
        port_list[pid] = pname
    return


con_dict_in = {}
con_cid = 5
am_cid = 10
def read_contract_info(filename, con_dict_in):
    wb = openpyxl.load_workbook(filename)
    ws = wb.active
    for r in ws.iter_rows(min_row=2, max_col=12, values_only=True):
        con_id = str(r[con_cid])[0:-3]
        con_am = float(r[am_cid])
        if con_id not in con_dict_in:
            con_dict_in[con_id] = con_am
        else:
            con_dict_in[con_id] += con_am


def regulate_port_str(in_str, port_list):
    if len(in_str) < 4:
        return ""
    port_id = in_str[:4]
    if port_id in port_list:
        return "{}({})".format(port_id, port_list[port_id])
    return ""


def regulate_con_id_str(in_str):
    if len(in_str) < 5:
        return ""
    if not (in_str[0].isalpha() and in_str[1].isalpha()
            and in_str[2].isdigit() and in_str[3].isdigit()):
        return ""
    return in_str


def rect_shrink(rect, v_step, h_step):
    res = (rect[0] + h_step, rect[1] + v_step,
           rect[2] - h_step, rect[3] - v_step)
    return res



def detect_text(img):
    return to.image_to_text(img, psm=to.PSM.SINGLE_LINE).rstrip()

def detect_chi_text(img):
    return to.image_to_text(img, lang="chi_sim+eng", psm=to.PSM.SINGLE_LINE).rstrip()


def denoise_by_components(in_img, min_size):
    img = in_img
    grey = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    bw = cv2.adaptiveThreshold(cv2.bitwise_not(grey), 255,
                               cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 25, -2)
    nb_components, output, stats, centroids = cv2.connectedComponentsWithStats(bw, connectivity=4)
    sizes = stats[1:, -1]
    nb_components = nb_components - 1
    #for every component in the image, you keep it only if it's above min_size
    for i in range(0, nb_components):
        if sizes[i] <= min_size:
            img[output == i + 1] = 255
    return img


## @brief recognize text from regions defined by ocr_table
def build_tex_item(org, iname, tables):
    t0 = tables[0]
    t1 = tables[1]
    if t0._h != 17 or t0._w != 2 or t1._h != 17 or t1._w != 2:
        return None
    file_str = "{}_{}_".format(iname.pdf, iname.idx)
    res = tex_item()
    res.pdf = iname.pdf
    res.pdf_idx = iname.idx
    img = image_preprocess(org)
    #img.save(os.path.join(str_output, file_str + "pre.png"))
    to = tables[0]._origin
    coner = tables[0]._pos[0]
    anchor = (to[0] + coner[0], to[1] + coner[1])
    date_rect = (anchor[0]-240, anchor[1] - 105, anchor[0] + 20, anchor[1] - 60)
    tmp_img = img.crop(date_rect)
    tmp_name = os.path.join(str_output, "tmp.png")
    tmp_img.save(tmp_name)
    tmp_cv = cv2.imread(tmp_name)
    tmp_cv = denoise_by_components(tmp_cv, 30)
    cv2.imwrite(tmp_name, tmp_cv)
    res.img_date = Image.open(tmp_name)
    #res.img_date = img.crop(date_rect)
    
    res.img_no = img.crop(rect_shrink(tables[0].rect(0, 0), 6, 3))
    res.img_idx = img.crop(rect_shrink(tables[1].rect(0, 0), 6, 3))
    res.img_tex_type = img.crop(rect_shrink(tables[0].rect(1, 0), 6, 3))
    res.img_port = img.crop(rect_shrink(tables[0].rect(2, 0), 6, 3))
    res.img_amount = img.crop(rect_shrink(tables[0].rect(9, 0), 6, 3))
    res.img_con_id = img.crop(rect_shrink(tables[1].rect(5, 0), 6, 3))
    try:    
        res.date = regulate_date_str(detect_chi_text(res.img_date))
        res.no = detect_text(res.img_no)
        res.idx = detect_text(res.img_idx).replace('o', '0').replace('O', '0').replace('a', '0')
        res.tex_type = regulate_type_str(detect_text(res.img_tex_type))
        res.port = regulate_port_str(detect_text(res.img_port), port_list)
        res.con_id = regulate_con_id_str(detect_text(res.img_con_id))
        str_am = detect_text(res.img_amount).replace(',', '.').replace('，', '.')
        res.amount = float(str_am)
        # try find amount error
        if str(res.amount) != str_am:
            res.amount = 0
    except:
        stem = os.path.splitext(os.path.basename(iname.path))[0]
        debug_dir = os.path.join(str_output, stem)
        if not os.path.exists(debug_dir):
            os.makedirs(debug_dir)

        res.img_date.save(os.path.join(debug_dir, file_str + "img_date.png"))
        res.img_no.save(os.path.join(debug_dir, file_str + "img_no.png"))
        res.img_idx.save(os.path.join(debug_dir, file_str + "img_idx.png"))
        res.img_tex_type.save(os.path.join(debug_dir, file_str + "img_tex_type.png"))
        res.img_port.save(os.path.join(debug_dir, file_str + "img_port.png"))
        res.img_amount.save(os.path.join(debug_dir, file_str + "img_amount.png"))
    return res


def image_preprocess(img):
    gray = img.convert('L')
    blackwhite = gray.point(lambda x: 0 if x < 200 else 255, '1')
    return blackwhite


def get_line_length(l):
    w = abs(l[0][2] - l[0][0]) 
    h = abs(l[0][1] - l[0][3])
    return w * w + h * h


def find_max_line(img_cv):
    gray = cv2.cvtColor(img_cv,cv2.COLOR_BGR2GRAY)
    kernel_size = 5
    h_pattern = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    v_pattern = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 3))
    gray = cv2.dilate(gray, h_pattern)
    dst = cv2.Canny(gray, 50, 200, 3);
    dst = cv2.dilate(dst, v_pattern)
    lines = cv2.HoughLinesP(dst, 1, 3.1415926 / 720, 100, 0, 0)
    if lines is None or len(lines) < 1:
        print("no line found")
        return None
    max_l = lines[0]
    max_len = 0
    for l in lines:
        cur_len = get_line_length(l)
        if cur_len > max_len:
            max_l = l
            max_len = cur_len
    return max_l


def rotate_horizontal(img, l):
    h, w = img.shape[:2]
    angle = math.atan((l[0][3] - l[0][1]) / (l[0][2] - l[0][0]))
    angle_d = angle / math.pi * 180
    print("image rotate {}".format(angle_d))
    mat = cv2.getRotationMatrix2D((h/2, w/2), angle_d, 1)
    res = cv2.warpAffine(img, mat, (w, h))
    # 转化角度为弧度
    theta = abs(angle)
    # 计算高宽比
    hw_ratio = float(h) / float(w)
    # 计算裁剪边长系数的分子项
    tan_theta = np.tan(theta)
    numerator = np.cos(theta) + np.sin(theta) * np.tan(theta)
    # 计算分母中和高宽比相关的项
    r = hw_ratio if h > w else 1 / hw_ratio
    # 计算分母项
    denominator = r * tan_theta + 1
    # 最终的边长系数
    crop_mult = numerator / denominator
    # 得到裁剪区域
    w_crop = int(crop_mult * w)
    h_crop = int(crop_mult * h)
    x0 = int((w - w_crop) / 2)
    y0 = int((h - h_crop) / 2)
    print("crop image {} {}".format(w_crop, h_crop))
    return res[y0 : y0 + h_crop, x0 : x0 + w_crop]


# find in first page
def get_total_from_pic(iname, con_id_list, am_list):
    con_id_list.clear()
    am_list.clear()
    tables = ocr_detect_table(iname)


def get_info_from_pic(iname, info_list):
    img_path = iname.path
    error_str = "{}_{}_".format(iname.pdf, iname.idx)
    # rotate image
    img_cv = cv2.imread(img_path)
    l = find_max_line(img_cv)
    if l is not None:
        img_rot = rotate_horizontal(img_cv, l)
        cv2.imwrite(img_path, img_rot)
    
    tables = ocr_detect_table(iname)
    #exclude invalid table contour
    ti = 0
    while ti < len(tables):
        t = tables[ti]
        if t._h != 17 or t._w != 2 or len(t._pos) != 34:
            tables.pop(ti)
        else:
            ti += 1
            
    img = Image.open(img_path)
    if len(tables) != 2:
        error_name = os.path.join(str_output, error_str + "error.png")
        img.save(error_name)
        print("Error! Fail to detect table from pdf {}, table number: {}".format(error_name, len(tables)))
        return
    item_res = build_tex_item(img, iname, tables)
    if item_res is not None:
        info_list.append(item_res)
    else:
        error_name = os.path.join(str_output, error_str + "error.png")
        img.save(error_name)
        print("Error! Fail to detect table from pdf {}".format(error_name))
        print("table0: {} x {}".format(tables[0]._w, tables[0]._h))
        print("table1: {} x {}".format(tables[1]._w, tables[1]._h))
        return

def get_info_from_pdf(pdf, info_list):
    images = convert_from_path(pdf)
    stem , _ = os.path.splitext(os.path.basename(pdf))
    for idx, img in enumerate(images):
        i_path = os.path.join(str_output, "{}_{}_img.png".format(stem, idx))
        img.save(i_path)
        iname = img_name(i_path, stem, idx)
        get_info_from_pic(iname, info_list)

def write_item_to_xls(filename, out_list):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["日期", "报关单号", "税单序号", "税种", "申报口岸", "税款金额", "文件位置"])
    for vi in out_list:
        ws.append([vi.date, vi.no, vi.idx, vi.tex_type, vi.port, vi.amount, "{}/{}".format(vi.pdf, vi.pdf_idx)])
    wb.save(filename)


def doc_add_cell_pic(cell, pic):
    img_name = os.path.join(str_output, "img_{}.png".format(3))
    arr = np.bitwise_not(np.asarray(pic))
    ind = np.nonzero(arr.any(axis=0))[0] # indices of non empty columns
    width = pic.width - 1
    if len(ind) > 0:
        width = int((ind[-1] - ind[0] + 1) * 1.2)
    rect = (0, 0, width, pic.height - 1)
    pic.crop(rect).save(img_name)
    if os.path.exists(img_name):
        pg = cell.paragraphs[0]
        run = pg.add_run()
        run.add_picture(img_name, height=Mm(5))


def set_cell_text(cell, text):
    cell.text = text
    if text == "":
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="AA0000"/>'.format(nsdecls('w'))))


def write_item_to_doc(filename, out_list):
    doc = docx.Document()
    i_num = len(out_list)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    row1 = table.rows[0]
    row1.cells[0].text = "日期"
    row1.cells[1].text = "报关单号"
    row1.cells[2].text = "税单序号"
    row1.cells[3].text = "税种"
    row1.cells[4].text = "申报口岸"
    row1.cells[5].text = "税款金额"
    row1.cells[6].text = "合同号"
    row1.cells[7].text = "文件位置"
    for idx, item in enumerate(out_list):
        txt_cells = table.add_row().cells
        set_cell_text(txt_cells[0], item.date)
        set_cell_text(txt_cells[1], item.no)
        if item.idx != "01" and item.idx != "02":
            item.idx = ""
        set_cell_text(txt_cells[2], item.idx)
        set_cell_text(txt_cells[3], item.tex_type)
        set_cell_text(txt_cells[4], item.port)
        # find error numbers
        str_am = str(item.amount)
        if str_am == "0" or str_am == "0.0" or str_am[-2:] == ".0":
            str_am = ""
        set_cell_text(txt_cells[5], str_am)
        set_cell_text(txt_cells[6], item.con_id)
        set_cell_text(txt_cells[7], "{}/{}".format(item.pdf, item.pdf_idx))
        pic_cells = table.add_row().cells
        doc_add_cell_pic(pic_cells[0], item.img_date)
        doc_add_cell_pic(pic_cells[1], item.img_no)
        doc_add_cell_pic(pic_cells[2], item.img_idx)
        doc_add_cell_pic(pic_cells[3], item.img_tex_type)
        doc_add_cell_pic(pic_cells[4], item.img_port)
        doc_add_cell_pic(pic_cells[5], item.img_amount)
        doc_add_cell_pic(pic_cells[6], item.img_con_id)
    doc.save(filename)


item_list = []


read_port_list(os.path.join(str_input, "duty.xlsx"), port_list)

dir_input = os.fsencode(str_input)
for file in os.listdir(dir_input):
    if not os.path.isfile(os.path.join(dir_input, file)):
        continue
    filename = os.fsdecode(file)
    stem, ext = os.path.splitext(filename)
    if ext == ".pdf":
        get_info_from_pdf(os.path.join(str_input, filename), item_list)
    elif ext == ".jpg" or ext == ".png":
        get_info_from_pic(img_name(os.path.join(str_input, filename), stem, 0), item_list)
    
#write_item_to_xls(os.path.join(str_output, "res.xlsx"), item_list)
write_item_to_doc(os.path.join(str_output, "res.docx"), item_list)
