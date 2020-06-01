# -*- coding: utf-8 -*-
## @file generate_docx.py
## @brief create documents from given folder structure
## @note generated py file meed pv module/framework_util to be parsed
## @author jiayanming

import os.path
import math
import sys
import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Mm
from test_framework.utils import g_config

## hausdorff relative#########################
class HausdorffSts:
    def __init__(self):
        self.sigma_rate = [0.0, 0.0, 0.0, 0.0, 0.0, 0.0]
        self.sigma_num = [0, 0, 0, 0, 0, 0]
        self.mean_total = 0.0
        self.mean_positive = 0.0
        self.mean_negtive = 0.0
        self.max_positive = 0.0
        self.max_negtive = 0.0
        self.standard_deviation = 0.0
        self.in_file = ""
        self.v_num = 0
        self.nominal_num = 0
        self.critical_num = 0
        self.max_num = 0
        self.out_num = 0
        self.nominal_dist = 0.0
        self.critical_dist = 0.0
        self.max_dist = 0.0

    def get_from_hd(self, sd):
        self.nominal_dist = float(g_config.config_val("hd_nominal_dist", "0.03"))
        self.critical_dist = float(g_config.config_val("hd_critical_dist", "0.05"))
        self.max_dist = float(g_config.config_val("hd_max_dist", "0.3"))

        fd = sd.GetFieldData()
        vtk_arr = fd.GetArray("six_sigma_rate")
        if vtk_arr is None or vtk_arr.GetDataSize() != 6:
            print("Warning! no statistics info in hausdorff output")
            return
        self.sigma_rate = [vtk_arr.GetTuple1(i) for i in range(0, 6)]
        vtk_arr = fd.GetArray("six_sigma_num")
        self.sigma_num = [vtk_arr.GetTuple1(i) for i in range(0, 6)]
        self.mean_total = fd.GetArray("mean_total").GetTuple1(0)
        self.mean_positive = fd.GetArray("mean_positive").GetTuple1(0)
        self.mean_negative = fd.GetArray("mean_negative").GetTuple1(0)
        self.max_positive = fd.GetArray("max_positive").GetTuple1(0)
        self.max_negative = fd.GetArray("max_negative").GetTuple1(0)
        self.standard_deviation = fd.GetArray("standard_deviation").GetTuple1(0)

        # calculate dist rate
        d_arr = sd.GetPointData().GetArray("Distance")
        self.v_num = d_arr.GetNumberOfTuples();
        nominal_num = 0
        critical_num = 0
        max_num = 0
        out_num = 0
        for vi in range(0, self.v_num):
            val = abs(d_arr.GetTuple1(vi))
            if val < self.nominal_dist:
                nominal_num += 1
            elif val < self.critical_dist:
                critical_num += 1
            elif val < self.max_dist:
                max_num += 1
            else:
                out_num += 1
        self.nominal_num = nominal_num
        self.critical_num = critical_num
        self.max_num = max_num
        self.out_nm = out_num

    def write_to_file(self, filename):
        f_sts = open(filename, "w", encoding='utf-8')
        f_sts.write("{}\n".format(" ".join(map(str, [self.sigma_rate.GetTuple1(i) for i in range(0, 6)]))))
        f_sts.write("{}\n".format(" ".join(map(str, [self.sigma_num.GetTuple1(i) for i in range(0, 6)]))))
        f_sts.write("{}\n".format(self.mean_total))
        f_sts.write("{}\n".format(self.mean_positive))
        f_sts.write("{}\n".format(self.mean_negative))
        f_sts.write("{}\n".format(self.max_positive))
        f_sts.write("{}\n".format(self.max_negative))
        f_sts.write("{}\n".format(self.standard_deviation))
        f_sts.write("{}\n".format(self.in_file))
        f_sts.write("{}\n".format(self.v_num))
        f_sts.write("{}\n".format(self.nominal_num))
        f_sts.write("{}\n".format(self.critical_num))
        f_sts.write("{}\n".format(self.max_num))
        f_sts.write("{}\n".format(self.out_num))
        f_sts.write("{}\n".format(self.nominal_dist))
        f_sts.write("{}\n".format(self.critical_dist))
        f_sts.write("{}\n".format(self.max_dist))
        f_sts.close()


    def read_from_file(self, filename):
        content = None
        if not os.path.exists(filename):
            print("Warning! no dist sts file found in {}".format(filename))
            return
        with open(filename, "r", encoding='utf-8') as f:
            content = f.readlines()
        if len(content) < 17:
            return
        str_list = [l.strip() for l in content]
        l_rate = str_list[0].split(" ")
        l_num = str_list[1].split(" ")
        self.sigma_rate.clear()
        self.sigma_num.clear()
        for item in l_rate:
            self.sigma_rate.append(float(item))
        for item in l_num:
            self.sigma_num.append(int(float(item)))
        self.mean_total = float(str_list[2])
        self.mean_positive = float(str_list[3])
        self.mean_negtive = float(str_list[4])
        self.max_positive = float(str_list[5])
        self.max_negtive = float(str_list[6])
        self.standard_deviation = float(str_list[7])
        self.in_file = str_list[8]
        self.v_num = int(str_list[9])
        self.nominal_num = float(str_list[10])
        self.critical_num = float(str_list[11])
        self.max_num = float(str_list[12])
        self.out_num = float(str_list[13])
        self.nominal_dist = float(str_list[14])
        self.critical_dist = float(str_list[15])
        self.max_dist = float(str_list[16])


## end hausdorff relative#######################


## @brief read camera positions in input config
def read_cam(case_file):
    if not os.path.exists(case_file):
        print("Warning! config file {} does not exists!".format(case_file))
        return []
    content = []
    with open(case_file, encoding="utf-8") as f:
        content = f.readlines()
    str_list = [l.strip() for l in content if len(l) > 20]
    str_lines = [line.split(", ") for line in str_list]
    res = [[float(s) for s in item] for item in str_lines if len(item) == 12]
    if len(res) == 0:
        print("Warning! No Camera Record in file {}!".format(case_file))
    return res


def add_cell_content(cell, text, pic):
    if os.path.exists(pic):
        pg = cell.paragraphs[0]
        run = pg.add_run()
        run.add_picture(pic, cell.width)
    else:
        print("Warning: No ScreenShot {}".format(pic))
    cell.add_paragraph(text)


def shade_cell(cell):
    shading_ele = parse_xml(r'<w:shd {} w:fill="B8B8B8"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_ele)


def merge_row(row, txt):
    cells = row.cells
    cells[0].merge(cells[len(cells) - 1])
    cells[0].text = txt
    return cells[0]



class DocxGenerator:
    def __init__(self, dir_input, dir_output, list_case, list_ver, list_alg):
        self._dirInput = dir_input
        self._dirOutput = dir_output
        self._listCase = list_case
        self._listVer = list_ver
        self._listAlg = list_alg
        self._doc = Document()

    # generate paraview project for given data
    def get_paraview_project(self, filename, case, l_ver, l_alg, compare_type, has_input):
        str_list_v = "["
        for v in l_ver:
            str_list_v = str_list_v + "\"{}\",".format(v)
        str_list_v = str_list_v[:-1] + "]"

        str_list_a = "["
        for a in l_alg:
            str_list_a = str_list_a + "\"{}\",".format(a)
        str_list_a = str_list_a[:-1] + "]"

        str_has_input = "True" if has_input else "False"
        old_ver = g_config.config_val("exe_old_pv_state", True)
        file_content = ""
        if old_ver:
            file_content = """# -*- coding: utf-8 -*-\n## @brief Paraview Macro to reproduce data state\n## @author jiayanming_auto_generate\nimport os\nimport sys\ndir_py_module = os.path.join(os.getcwd(), \"..\", \"Sn3D_plugins\", \"scripts\", \"pv_module\")\nsys.path.append(dir_py_module)\nfrom framework_util import *\nload_state_files_v1(r\"{}\", r\"{}\", \"{}\", {}, {}, \"{}\", {})\n""".format(self._dirInput, self._dirOutput, case, str_list_v, str_list_a, compare_type, str_has_input)
        else:
            file_content = """# -*- coding: utf-8 -*-\n## @brief Paraview Macro to reproduce data state\n## @author jiayanming_auto_generate\nimport os\nimport sys\nfrom test_framework.framework_util import load_state_files_v1\nload_state_files_v1(r\"{}\", r\"{}\", \"{}\", {}, {}, \"{}\", {})\n""".format(self._dirInput, self._dirOutput, case, str_list_v, str_list_a, compare_type, str_has_input)
        with open(filename, "w", encoding="utf-8") as text_file:
            text_file.write(file_content)

    def add_case_table_ver(self, case, cam_num):
        dir_state = os.path.join(self._dirOutput, "ParaView_projects")
        if not os.path.exists(dir_state):
            os.makedirs(dir_state)

        self._doc.add_paragraph("Compare between Versions")
        col_num = len(self._listVer)
        if col_num < 1:
            print("Error: No Version Checked!")
            return 1
        str_time = str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
        has_input = "input" in self._listVer
        for alg in self._listAlg:
            name_state = "{}_{}_{}.py".format(case.replace("/", "_"), alg, str_time)
            file_state = os.path.join(dir_state, name_state)
            self.get_paraview_project(file_state, case,  self._listVer, [alg], "Versions", has_input)
            table = self._doc.add_table(rows=1, cols=col_num)
            table.style = 'Table Grid'
            # color title
            shade_cell(merge_row(table.rows[0], alg))
            shade_cell(merge_row(table.add_row(), file_state))
            # add row for every camera angle
            for cam in range(0, cam_num):
                for vi, cell in enumerate(table.add_row().cells):
                    ver = self._listVer[vi]
                    dir_pic = os.path.join(self._dirOutput, case, ver)
                    name_pic = "ss_{}_v{}.png".format(alg, cam)
                    if ver == "input":
                        name_pic = "ss_input_v{}.png".format(cam)
                    file_pic = os.path.join(dir_pic, name_pic)
                    add_cell_content(cell, ver, file_pic)
        return 0

    def add_case_table_alg(self, case, cam_num):
        if len(self._listAlg) < 1:
            print("Error: No FileName Checked!")
            return 1

        dir_state = os.path.join(self._dirOutput, "ParaView_projects")
        if not os.path.exists(dir_state):
            os.makedirs(dir_state)
        self._doc.add_paragraph("Compare between Files")
        compare_with_input = False
        col_list = self._listAlg.copy()
        if "input" in self._listVer:
            compare_with_input = True
            col_list.insert(0, "input")
        str_time = str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
        for ver in self._listVer:
            if ver == "input": # skip input table, because no file to compare in input
                continue
            name_state = "{}_{}_{}.py".format(case.replace("/", "_"), ver, str_time)
            file_state = os.path.join(dir_state, name_state)
            self.get_paraview_project(file_state, case, [ver], self._listAlg, "FileNames", compare_with_input)
            table = self._doc.add_table(rows=1, cols=len(col_list))
            table.style = 'Table Grid'
            # color title
            shade_cell(merge_row(table.rows[0], ver))
            shade_cell(merge_row(table.add_row(), file_state))
            # add row for every camera angle
            for cam in range(0, cam_num):
                for fi, cell in enumerate(table.add_row().cells):
                    alg = col_list[fi]
                    dir_pic = os.path.join(self._dirOutput, case, ver)
                    name_pic = "ss_{}_v{}.png".format(alg, cam)
                    if alg == "input":
                        dir_pic = os.path.join(self._dirOutput, case, "input")
                    file_pic = os.path.join(dir_pic, name_pic)
                    add_cell_content(cell, alg, file_pic)
        return 0

    def add_hausdorff_case_table(self, case, cam_num):
        col_num = len(self._listVer)
        if col_num < 1:
            print("Error: No Version Checked!")
            return 1
        for alg in self._listAlg:
            table = self._doc.add_table(rows=1, cols=col_num)
            table.style = 'Table Grid'
            row1 = table.rows[0]
            row1.cells[0].merge(row1.cells[col_num - 1])
            row1.cells[0].text = alg
            case_tmp = case.replace("/", "_")
            name_state = "{}_{}_{}.py".format(case_tmp, alg, str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S")))
            # color title
            shade_cell(row1.cells[0])
            if g_config.config_val("ss_default_camera_type", "4_quadrant") == "4_quadrant":
                cam_num = cam_num + 4
            else:
                cam_num = cam_num + 6
            for cam in range(0, cam_num):
                row_cells = table.add_row().cells
                for vi in range(0, col_num):
                    ver = self._listVer[vi]
                    dir_pic = os.path.join(self._dirOutput, case, ver)
                    name_pic = "ss_{}_v{}.png".format(alg, cam)
                    if ver == "input":
                        name_pic = "ss_input_v{}.png".format(cam)
                    file_pic = os.path.join(dir_pic, name_pic)
                    add_cell_content(row_cells[vi], ver, file_pic)
        return 0
        

    def fill_row(self, table, rid, content):
        r = table.rows[rid].cells
        shade_cell(r[0])
        for ci in range(0, len(content)):
            if isinstance(content[ci], float):
                r[ci].text = "{:.5f}".format(content[ci])
            else:
                r[ci].text = str(content[ci])



    def add_hausdorff_statistic_table(self, case):
        if len(self._listVer) != 2:
            return
        # read sts info
        dir_a2b = os.path.join(self._dirOutput, case, self._listVer[0])
        dir_b2a = os.path.join(self._dirOutput, case, self._listVer[1])
        a2b = HausdorffSts()
        b2a = HausdorffSts()
        a2b.read_from_file(os.path.join(dir_a2b, "dist.sts"))
        b2a.read_from_file(os.path.join(dir_b2a, "dist.sts"))
        add_hausdorff_statistic_table(self, a2b, b2a)


    def add_hausdorff_statistic_table(self, a2b, b2a):
        # heading
        self._doc.add_paragraph("Deviation Report between two mesh A and B")
        self._doc.add_paragraph("A: {}".format(a2b.in_file))
        self._doc.add_paragraph("B: {}".format(b2a.in_file))
        # build table
        if g_config.config_val("hd_standard_statistics", True):
            self._doc.add_paragraph("")
            self._doc.add_paragraph("General Statistics")
            table = self._doc.add_table(rows = 7, cols = 3)
            table.style = 'Table Grid'
            self.fill_row(table, 0, ["", "A to B", "B to A"])
            shade_cell(table.rows[0].cells[1])
            shade_cell(table.rows[0].cells[2])
            self.fill_row(table, 1, ["mean_total", a2b.mean_total, b2a.mean_total])
            self.fill_row(table, 2, ["standard_deviation", a2b.standard_deviation, b2a.standard_deviation])
            self.fill_row(table, 3, ["mean_positive", a2b.mean_positive, b2a.mean_positive])
            self.fill_row(table, 4, ["mean_negtive", a2b.mean_negtive, b2a.mean_negtive])
            self.fill_row(table, 5, ["max_positive", a2b.max_positive, b2a.max_positive])
            self.fill_row(table, 6, ["max_negtive", a2b.max_negtive, b2a.max_negtive])
        
        if g_config.config_val("hd_distance_rate", True):
            self._doc.add_paragraph("")
            self._doc.add_paragraph("Distance Percentage Statistics A to B")
            t_rate = self._doc.add_table(rows = 5, cols = 3)
            t_rate.style = 'Table Grid'
            self.fill_row(t_rate, 0, ["", "A to B"])
            shade_cell(t_rate.rows[0].cells[1])
            shade_cell(t_rate.rows[0].cells[2])
            nominal_num = int(a2b.nominal_num)
            critical_num = int(nominal_num + a2b.critical_num)
            max_num = int(critical_num + a2b.max_num)
            out_num = int(a2b.out_num)
            v_num = float(a2b.v_num)
            self.fill_row(t_rate, 0, ["", "Point Number", "Point Percentage"])
            self.fill_row(t_rate, 1, ["Nominal(<{})".format(a2b.nominal_dist),
                                      nominal_num,
                                      "{:.2f}%".format(float(nominal_num) / v_num * 100)])
            self.fill_row(t_rate, 2, ["Critical(<{})".format(a2b.critical_dist),
                                      critical_num,
                                      "{:.2f}%".format(float(critical_num) / v_num * 100)])
            self.fill_row(t_rate, 3, ["Max(<{})".format(a2b.max_dist), max_num,
                                      "{:.2f}%".format(float(max_num) / v_num * 100)])
            self.fill_row(t_rate, 4, ["Out(>={})".format(a2b.max_dist), out_num,
                                      "{:.2f}%".format(float(out_num) / v_num * 100)])

            self._doc.add_paragraph("")
            self._doc.add_paragraph("Distance Percentage Statistics B to A")
            t_rate_b2a = self._doc.add_table(rows = 5, cols = 3)
            t_rate_b2a.style = 'Table Grid'
            self.fill_row(t_rate_b2a, 0, ["", "B to A"])
            shade_cell(t_rate_b2a.rows[0].cells[1])
            shade_cell(t_rate_b2a.rows[0].cells[2])
            nominal_num = int(b2a.nominal_num)
            critical_num = int(nominal_num + b2a.critical_num)
            max_num = int(critical_num + b2a.max_num)
            out_num = int(b2a.out_num)
            v_num = float(b2a.v_num)
            self.fill_row(t_rate_b2a, 0, ["", "Point Number", "Point Percentage"])
            self.fill_row(t_rate_b2a, 1, ["Nominal(<{})".format(b2a.nominal_dist),
                                          nominal_num,
                                      "{:.2f}%".format(float(nominal_num) / v_num * 100)])
            self.fill_row(t_rate_b2a, 2, ["Critical(<{})".format(b2a.critical_dist),
                                          critical_num,
                                      "{:.2f}%".format(float(critical_num) / v_num * 100)])
            self.fill_row(t_rate_b2a, 3, ["Max(<{})".format(b2a.max_dist), max_num,
                                      "{:.2f}%".format(float(max_num) / v_num * 100)])
            self.fill_row(t_rate_b2a, 4, ["Out(>={})".format(b2a.max_dist), out_num,
                                      "{:.2f}%".format(float(out_num) / v_num * 100)])

        if g_config.config_val("hd_6_sigma", True):
            self._doc.add_paragraph("")
            self._doc.add_paragraph("6-SIGMA Statistics A to B")
            ts_a2b = self._doc.add_table(rows = 7, cols = 4)
            ts_a2b.style = 'Table Grid'
            self.fill_row(ts_a2b, 0, ["", "Point Number", "Point Percentage"])
            shade_cell(ts_a2b.rows[0].cells[1])
            shade_cell(ts_a2b.rows[0].cells[2])
            sigma_map = [-3, -2, -1, 1, 2, 3]
            for ri in range(0, 6):
                self.fill_row(ts_a2b, ri + 1, ["{} sigma".format(sigma_map[ri]),
                                               a2b.sigma_num[ri], "{:.2f}%".format(a2b.sigma_rate[ri] * 100.0)])
            # histogram
            ts_a2b.rows[0].cells[3].merge(ts_a2b.rows[6].cells[3])

            self._doc.add_paragraph("")
            self._doc.add_paragraph("6-SIGMA Statistics B to A")
            ts_b2a = self._doc.add_table(rows = 7, cols = 4)
            ts_b2a.style = 'Table Grid'
            self.fill_row(ts_b2a, 0, ["", "Point Number", "Point Percentage"])
            shade_cell(ts_b2a.rows[0].cells[1])
            shade_cell(ts_b2a.rows[0].cells[2])
            for ri in range(0, 6):
                self.fill_row(ts_b2a, ri + 1, ["{} sigma".format(sigma_map[ri]),
                                           b2a.sigma_num[ri], "{:.2f}%".format(b2a.sigma_rate[ri] * 100.0)])
            ts_b2a.rows[0].cells[3].merge(ts_b2a.rows[6].cells[3])        

    def add_hausdorff_statistic_table_single(self, case):
        if len(self._listVer) != 1:
            return
        # read sts info
        dir_a2b = os.path.join(self._dirOutput, case, self._listVer[0])
        dir_b2a = os.path.join(self._dirOutput, case, "hausdorff_B2A")
        a2b = HausdorffSts()
        b2a = HausdorffSts()
        a2b.read_from_file(os.path.join(dir_a2b, "dist.sts"))
        b2a.read_from_file(os.path.join(dir_b2a, "dist.sts"))
        # heading
        self._doc.add_paragraph("Deviation Report between two mesh A and B")
        self._doc.add_paragraph("A: {}".format(a2b.in_file))
        self._doc.add_paragraph("B: {}".format(b2a.in_file))
        if g_config.config_val("hd_standard_statistics", True):
            self._doc.add_paragraph("")
            self._doc.add_paragraph("General Statistics")
            # build table
            table = self._doc.add_table(rows = 7, cols = 2)
            table.style = 'Table Grid'
            self.fill_row(table, 0, ["", "A to B"])
            shade_cell(table.rows[0].cells[1])
            self.fill_row(table, 1, ["mean_total", a2b.mean_total])
            self.fill_row(table, 2, ["standard_deviation", a2b.standard_deviation])
            self.fill_row(table, 3, ["mean_positive", a2b.mean_positive])
            self.fill_row(table, 4, ["mean_negtive", a2b.mean_negtive])
            self.fill_row(table, 5, ["max_positive", a2b.max_positive])
            self.fill_row(table, 6, ["max_negtive", a2b.max_negtive])

        if g_config.config_val("hd_distance_rate", True):
            self._doc.add_paragraph("")
            self._doc.add_paragraph("Distance Percentage Statistics A to B")
            t_rate = self._doc.add_table(rows = 5, cols = 3)
            t_rate.style = 'Table Grid'
            self.fill_row(t_rate, 0, ["", "A to B"])
            shade_cell(t_rate.rows[0].cells[1])
            shade_cell(t_rate.rows[0].cells[2])
            nominal_num = int(a2b.nominal_num)
            critical_num = int(nominal_num + a2b.critical_num)
            max_num = int(critical_num + a2b.max_num)
            out_num = int(a2b.out_num)
            v_num = float(a2b.v_num)
            self.fill_row(t_rate, 0, ["", "Point Number", "Point Percentage"])
            self.fill_row(t_rate, 1, ["Nominal(<{})".format(a2b.nominal_dist), nominal_num,
                                      "{:.2f}%".format(float(nominal_num) / v_num * 100)])
            self.fill_row(t_rate, 2, ["Critical(<{})".format(a2b.critical_dist), critical_num,
                                      "{:.2f}%".format(float(critical_num) / v_num * 100)])
            self.fill_row(t_rate, 3, ["Max(<{})".format(a2b.max_dist), max_num,
                                      "{:.2f}%".format(float(max_num) / v_num * 100)])
            self.fill_row(t_rate, 4, ["Out(>={})".format(a2b.max_dist), out_num,
                                      "{:.2f}%".format(float(out_num) / v_num * 100)])

        if g_config.config_val("hd_6_sigma", True):
            self._doc.add_paragraph("")
            self._doc.add_paragraph("6-SIGMA Statistics A to B")
            ts_a2b = self._doc.add_table(rows = 7, cols = 4)
            ts_a2b.style = 'Table Grid'
            self.fill_row(ts_a2b, 0, ["", "Point Number", "Point Percentage"])
            shade_cell(ts_a2b.rows[0].cells[1])
            shade_cell(ts_a2b.rows[0].cells[2])
            sigma_map = [-3, -2, -1, 1, 2, 3]
            for ri in range(0, 6):
                self.fill_row(ts_a2b, ri + 1, ["{} sigma".format(sigma_map[ri]),
                                               a2b.sigma_num[ri], "{:.2f}%".format(a2b.sigma_rate[ri] * 100.0)])
            # histogram
            ts_a2b.rows[0].cells[3].merge(ts_a2b.rows[6].cells[3])


    ## @brief generate docx file from given algorithm output dir, and config
    ## @param dir_input data case config directory
    ## @param dir_output algorithm/screenshots output directory
    ## @param file_config file contains user specified compare config
    def generate_docx(self, file_save, file_config):
        compare_type = g_config.config_val("doc_compare_type", "Versions")
        doc = self._doc
        # screen shot view list
        str_time = str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
        doc.add_heading("Compare Result {}".format(str_time), 0)
        doc.add_paragraph("From config file: {}".format(file_config))
        for case in self._listCase:
            doc.add_paragraph("")
            doc.add_paragraph(case, style='List Bullet')
            list_cam = read_cam(os.path.join(self._dirInput, case, "config.txt"))
            res = 0
            if compare_type == "Versions":
                res = self.add_case_table_ver(case, len(list_cam))
            else:
                res = self.add_case_table_alg(case, len(list_cam))
            if res != 0:
                print("Case Table Error for case: {}".format(case))
        doc.save(file_save)

    def generate_hausdorff_docx(self, file_save, file_config):
        doc = self._doc
        # screen shot view list
        str_time = str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
        doc.add_heading("Compare Result {}".format(str_time), 0)
        doc.add_paragraph("From config file: {}".format(file_config))
        for case in self._listCase:
            doc.add_paragraph(case, style='List Bullet')
            list_cam = read_cam(os.path.join(self._dirInput, case, "config.txt"))
            if list_cam is None or len(list_cam) == 0:
                print("Warning! no cam table for {}".format(case))
                list_cam = []
            # hausdorff
            ver_num = len(self._listVer)
            if ver_num > 0 and self._listVer[0] == "hausdorff_A2B":
                if ver_num == 2:
                    self.add_hausdorff_statistic_table(case)
                else:
                    self.add_hausdorff_statistic_table_single(case)
                if not g_config.config_val("hd_screenshot_table", True):
                    continue
            doc.add_paragraph("")
            doc.add_paragraph("ScreenShots Compare Tables")
            if self.add_hausdorff_case_table(case, len(list_cam)) != 0:
                print("Case Table Error for case: {}".format(case))
        doc.save(file_save)


def generate_docx_wrap(dir_input, dir_output, file_config):
    ss = utils.SessionConfig()
    if not ss.read_config(file_config):
        return

    d = DocxGenerator(dir_input, dir_output, sc.list_case, sc.list_ver, sc.list_alg)
    # input data read from config file
    config_stem = os.path.splitext(os.path.split(file_config)[1])[0]
    str_time = str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
    file_save = os.path.join(dir_output, "{}_{}.docx".format(config_stem, str_time))
    if len(sys.argv) == 2:
        file_save = str(sys.argv[1])
    d.generate_docx(file_save)


if __name__ == "__main__":
    dir_input = "d:/data/test_framwork/input/"
    dir_output = "d:/data/test_framwork/output/"
    file_config = "d:/data/test_framwork/compare.txt"
    generate_docx_wrap(dir_input, dir_output, file_config)
