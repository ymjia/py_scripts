# -*- coding: utf-8 -*-
## @file generate_docx.py
## @brief create documents from given folder structure
## @note generated py file meed pv module/framework_util to be parsed
## @author jiayanming

import os.path
import math
import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Mm

## @brief read user concerned case name list
def read_config_list(config_str, pattern):
    lc = len(config_str)
    lp = len(pattern)
    if lc < lp:
        return None
    if config_str[0:lp] != pattern:
        return None
    return config_str[lp+1:].split(" ")



def read_compare_config(file_config):
    if not os.path.exists(file_config):
        return None
    case_list = []
    ver_list = []
    alg_list = []
    content = None
    with open(file_config) as f:
        content = f.readlines()
    str_list = [l.strip() for l in content]
    for line in str_list:
        if line[0:3] == "cas":
            case_list = read_config_list(line, "cas")
        elif line[0:3] == "ver":
            ver_list = read_config_list(line, "ver")
        elif line[0:3] == "alg":
            alg_list = read_config_list(line, "alg")
    return case_list, ver_list, alg_list


# generate paraview project for given data
def get_paraview_project(filename, dir_in, case, alg, list_v):
    str_list_v = "["
    for v in list_v:
        str_list_v = str_list_v + "\"{}\",".format(v)
    str_list_v = str_list_v[:-1] + "]"
    file_content = """# -*- coding: utf-8 -*-\n## @brief Paraview Macro to reproduce data state\n## @author jiayanming_auto_generate\nimport os\nimport sys\ndir_py_module = os.path.join(os.getcwd(), \"..\", \"Sn3D_plugins\", \"scripts\", \"pv_module\")\nsys.path.append(dir_py_module)\nfrom framework_util import *\nload_state_files(\"{}\", \"{}\", \"{}\", {})\n""".format(dir_in, case, alg, str_list_v)
    with open(filename, "w") as text_file:
        text_file.write(file_content)

def add_cell_content(cell, text, pic):
    pg = cell.paragraphs[0]
    run = pg.add_run()
    run.add_picture(pic, cell.width)
    cell.add_paragraph(text)


def add_case_table(doc, dir_in, case):
    for alg in list_alg:
        table = document.add_table(rows=1, cols=col_num)
        table.style = 'Table Grid'
        row1 = table.rows[0]
        row1.cells[0].merge(row1.cells[col_num - 1])
        row1.cells[0].text = alg
        row2 = table.add_row().cells
        row2[0].merge(row2[col_num-1])
        name_state = "{}_{}_{}.py".format(case, alg, str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S")))
        file_state = os.path.join(dir_in, name_state)
        row2[0].text = file_state
        get_paraview_project(file_state, dir_in, case, alg, list_ver)
        for cam in range(0, cam_num):
            row_cells = table.add_row().cells
            for vi in range(0, col_num):
                ver = list_ver[vi]
                dir_pic = os.path.join(dir_in, case, ver)
                name_pic = "ss_{}_v{}.png".format(alg, cam)
                file_pic = os.path.join(dir_pic, name_pic)
                add_cell_content(row_cells[vi], ver, file_pic)


dir_output = "d:/data/test_framwork/output/"
file_config = None

# input data read from config file
list_case, list_ver, list_alg = read_compare_config(file_config)
# screen shot view list
cam_num = 3

col_num = len(list_ver)
document = Document()

document.add_heading("compare", 0)
for case in list_case:
    document.add_paragraph(
    case, style='List Bullet')
    add_case_table(document, dir_output, case)


file_save = os.path.join(dir_output, "Compare_{}.docx".format(str(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))))
if len(sys.argv) == 2:
    file_save = str(sys.argv[1])
print(file_save)
document.save(file_save)
