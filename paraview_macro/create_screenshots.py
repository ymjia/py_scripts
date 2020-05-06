# -*- coding: utf-8 -*-
## @file create_screenshots.py
## @brief create screenshots for mesh in given position
## @author jiayanming

import os.path
import math
from paraview.simple import *
from paraview.simple import _active_objects

from docx import Document
from docx.shared import Inches


class ScreenShotHelper:
    def __init__(self):
        paraview.simple._DisableFirstRenderCameraReset()

    CameraPosition = [0, 0, 0]
    CameraFocalPoint = [0, 0, 0]
    CameraViewUp = [0, 0, 0]
    CameraParallelScale = 25.46296403946468
    def set_view_pos(self, view):
        view.CameraPosition = self.CameraPosition
        view.CameraFocalPoint = self.CameraFocalPoint
        view.CameraViewUp = self.CameraViewUp
        view.CameraParallelScale = self.CameraParallelScale

    def take_shot(self, view, source, filename):
        HideAll(view)
        Show(source, view)
        self.set_view_pos(view)
        view.Update()
        v_size = cur_v.ViewSize
        SaveScreenshot(filename, cur_v, ImageResolution=v_size, TransparentBackground=1)


#setup active object
active_obj = _active_objects()
cur_s = active_obj.get_source()
cur_v = active_obj.get_view()
s_list = active_obj.get_selected_sources()


pxm = servermanager.ProxyManager();
s_name = os.path.splitext(pxm.GetProxyName("sources", cur_s))[0]
ss = ScreenShotHelper()
ss.CameraPosition = cur_v.CameraPosition
ss.CameraFocalPoint = cur_v.CameraFocalPoint
ss.CameraViewUp = cur_v.CameraViewUp
ss.CameraParallelScale = cur_v.CameraParallelScale

# create a new 'SnDenoiseBilateral'
snDenoiseBilateral = SnDenoiseBilateral(Input=cur_s)
snDenoiseBilateral.NormalIteration = 2
RenameSource("{}_denoise".format(s_name), snDenoiseBilateral)
ss.take_shot(cur_v, snDenoiseBilateral, "D:/tmp/ss_denoise.png")

# create a new 'SnSmooth'
snFill = SnFillHole(Input=snDenoiseBilateral)
#Show(snFill, cur_v)
RenameSource("{}_smooth".format(s_name), snFill)
ss.take_shot(cur_v, snFill, 'D:/tmp/ss_smooth.png')

document = Document()

document.add_header("对比图", 0)
document.add_paragraph(
    'after denoise', style='List Bullet'
)

document.add_picture('d:/tmp/ss_denoise.png', width=Inches(1.25))
document.add_paragraph(
    'after smooth', style='List Number'
)
document.add_picture('d:/tmp/ss_denoise.png', width=Inches(2.25))
document.add_picture('d:/tmp/ss_smooth.png', width=Inches(2.25))
document.save('d:/tmp/compare.docx')

